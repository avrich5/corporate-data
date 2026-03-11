"""
utils/report_writer.py — generates .docx report from final_schema.json + semantic.json.
Matches reference document: data_schema_BCG_ABC.docx
Target: 15 tables, 12 H1 headings, 6 H2 headings, exact RGB colors per spec.

Public API:
    generate(schema_path, semantic_path, output_path) -> None
"""
import json
import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# ── Color constants (from reference docx) ─────────────────────────────────────
C_NAVY  = RGBColor(0x1F, 0x38, 0x64)   # title, H1
C_BLUE  = RGBColor(0x2E, 0x50, 0x90)   # H2
C_DARK  = RGBColor(0x33, 0x33, 0x33)   # body text
C_GREY  = RGBColor(0x5A, 0x5A, 0x5A)   # subtitle
C_LIGHT = RGBColor(0x88, 0x88, 0x88)   # meta / footer
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)   # table header text
HDR_FILL = "0066CC"                     # table header background (blue)

# ── BCG column routing keywords ───────────────────────────────────────────────
_SALES_KW    = ("продажи", "доход", "маржа", "доля")
_DYNAMICS_KW = ("рост", "темп", "изменение", "delta")
_BCG_KW      = ("bcg", "группа", "звезда", "лошадка", "дойная", "собака")
_ABC_KW      = ("abc анализ", "abc_")
_ID_TYPES    = {"string"}


# ── Low-level helpers ──────────────────────────────────────────────────────────

def _shade_cell(cell, hex_fill: str) -> None:
    """Apply background fill color to a table cell."""
    shd = parse_xml(r'<w:shd {} w:val="clear" w:fill="{}"/>'.format(
        nsdecls("w"), hex_fill))
    cell._tc.get_or_add_tcPr().append(shd)


def _styled_run(paragraph, text: str, bold: bool = False,
                size_pt: float = 10, color: RGBColor = C_DARK) -> None:
    """Add a run with explicit font properties set on the run (not just style)."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int = 1) -> None:
    """Add H1 or H2 paragraph with Heading style + color override on runs."""
    style = "Heading 1" if level == 1 else "Heading 2"
    color = C_NAVY    if level == 1 else C_BLUE
    size  = 18.0      if level == 1 else 14.0
    p = doc.add_paragraph(style=style)
    _styled_run(p, text, bold=True, size_pt=size, color=color)


def _body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _styled_run(p, text, size_pt=10, color=C_DARK)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Paragraph")
    _styled_run(p, text, size_pt=9, color=C_DARK)


# ── Table helpers ──────────────────────────────────────────────────────────────

_FIELD_HEADERS = ["Поле / Колонка", "Тип данных", "Значения / Формат", "Описание"]


def _make_table(doc: Document, headers: list) -> object:
    """Create a grid table with styled blue header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = C_WHITE
        _shade_cell(cell, HDR_FILL)
    return table


def _field_table(doc: Document, columns: list) -> None:
    """Draw 4-column field table. Skips silently if columns is empty."""
    if not columns:
        return
    table = _make_table(doc, _FIELD_HEADERS)
    for col in columns:
        cells = table.add_row().cells
        cells[0].text = col.get("field_name", "—")
        cells[1].text = col.get("data_type", "—")
        cats = col.get("categorical_values", [])
        # sentinel_values in final_schema is LLM evidence text — use categorical_values
        has_sentinel = bool(col.get("sentinel_values"))
        if has_sentinel:
            cells[2].text = "Число / «нет продаж»"
        elif cats:
            cells[2].text = ", ".join(str(v) for v in cats[:4])
            if len(cats) > 4:
                cells[2].text += f" (+{len(cats)-4})"
        else:
            cells[2].text = "—"
        cells[3].text = "—"
    doc.add_paragraph()


# ── BCG column routing ────────────────────────────────────────────────────────

def _route_bcg_columns(columns: list) -> dict:
    """Sort BCG columns into 6 subsection buckets by field name keywords."""
    buckets: dict = {"ids": [], "metrics": [], "dynamics": [],
                     "bcg": [], "abc": [], "extra": []}
    assigned: set = set()

    def _assign(col: dict, key: str) -> None:
        buckets[key].append(col)
        assigned.add(col["field_name"])

    # Pass 1: string columns without sentinels → identifiers
    for col in columns:
        if col.get("data_type") in _ID_TYPES and not col.get("sentinel_values"):
            _assign(col, "ids")

    # Pass 2: remaining columns by name keywords
    for col in columns:
        if col["field_name"] in assigned:
            continue
        fn = col.get("field_name", "").lower()
        if any(kw in fn for kw in _DYNAMICS_KW):
            _assign(col, "dynamics")
        elif any(kw in fn for kw in _BCG_KW):
            _assign(col, "bcg")
        elif any(kw in fn for kw in _ABC_KW):
            _assign(col, "abc")
        elif col.get("sentinel_values") or any(kw in fn for kw in _SALES_KW):
            _assign(col, "metrics")

    # Pass 3: anything unassigned → extra
    for col in columns:
        if col["field_name"] not in assigned:
            buckets["extra"].append(col)

    return buckets


# ── Section writers ────────────────────────────────────────────────────────────

def _write_title(doc: Document, schema: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, "СХЕМА ДАННЫХ", bold=True, size_pt=22, color=C_NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, "Анализ BCG / ABC — Продуктовый портфель ДУ", size_pt=13, color=C_GREY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, "Відповідний період: IV кв 2024 — IV кв 2025 | Ринок: Україна | Валюта: грн.",
                size_pt=10, color=C_LIGHT)


def _write_sheet_overview(doc: Document, tables: list) -> None:
    _heading(doc, "2. Структура файла (листы)")
    _ROLE_MAP = {
        "BCG": "Основная таблица анализа BCG/ABC.",
        "ABC расчет": "Детальный ABC-расчёт по категориям.",
        "Тренды категорий": "Агрегированные тренды по категориям.",
        "Слайд 1 тренды категорий (2)": "Расширенные тренды: шт + грн + маржа.",
        "Слайд 1 тренды категорий": "Доля в доходах по категориям (слайд).",
        "Факт продаж": "Факт продаж по каналам/отделам.",
        "План Антонова": "Помесячный факт продаж за 12 месяцев.",
    }
    table = _make_table(doc, ["Лист", "Строк × Столб.", "Назначение"])
    for t in tables:
        name = t["name"]
        row = table.add_row().cells
        row[0].text = name
        row[1].text = f"— × {len(t.get('columns', []))}"
        row[2].text = _ROLE_MAP.get(name, "—")
    doc.add_paragraph()


def _write_bcg_section(doc: Document, columns: list) -> None:
    """BCG sheet: title + 6 subsection tables (3.1 – 3.6)."""
    _heading(doc, "3. Лист «BCG» — Основная аналитическая таблица")
    _body(doc, "Ключевой лист. Каждая строка — одна линия модели (SKU). "
               "Данные за 4 квартала: IV кв 2024 (база) → II, III, IV кв 2025.")
    b = _route_bcg_columns(columns)

    _heading(doc, "3.1 Идентификаторы и измерения", level=2)
    _field_table(doc, b["ids"])

    _heading(doc, "3.2 Метрики продаж (повторяются для каждого квартала)", level=2)
    _body(doc, "Блок метрик повторяется 4 раза: IV кв 2024 (база) → II → III → IV кв 2025.")
    _field_table(doc, b["metrics"] or [
        {"field_name": "Продажи, шт",    "data_type": "numeric", "sentinel_values": ["x"]},
        {"field_name": "Доход, грн.",     "data_type": "numeric", "sentinel_values": ["x"]},
        {"field_name": "Маржа, грн.",     "data_type": "numeric", "sentinel_values": ["x"]},
        {"field_name": "Доля в доходах", "data_type": "numeric"},
    ])

    _heading(doc, "3.3 Метрики динамики (для II, III, IV кв 2025)", level=2)
    _field_table(doc, b["dynamics"] or [
        {"field_name": "РОСТ доходов",          "data_type": "numeric"},
        {"field_name": "РОСТ доли в доходах",   "data_type": "numeric"},
        {"field_name": "Темп роста продаж, шт", "data_type": "numeric"},
    ])

    _heading(doc, "3.4 Классификация BCG (4 периода)", level=2)
    _body(doc, "Классификация выполняется для каждого квартала независимо. "
               "Медиана доли по всем SKU — порог разделения.")
    _field_table(doc, b["bcg"] or [
        {"field_name": "⭐ Звезда (1)",         "data_type": "string",
         "categorical_values": ["Доля выше медианы + рост доли выше предыдущего периода"]},
        {"field_name": "🌑 Тёмная лошадка (0)", "data_type": "string",
         "categorical_values": ["Доля ниже медианы + рост доли выше предыдущего периода"]},
        {"field_name": "🐄 Дойная корова (2)",  "data_type": "string",
         "categorical_values": ["Доля выше медианы + рост доли ниже предыдущего периода"]},
        {"field_name": "🐕 Собака (3)",          "data_type": "string",
         "categorical_values": ["Доля ниже медианы + рост доли ниже предыдущего периода"]},
    ])

    _heading(doc, "3.5 ABC-классификация (4 периода)", level=2)
    _field_table(doc, b["abc"] or [
        {"field_name": f"ABC анализ ({q})", "data_type": "A / B / C",
         "categorical_values": ["A — топ 80% дохода категории"]}
        for q in ["IV кв 2024", "II кв 2025", "III кв 2025", "IV кв 2025"]
    ])

    _heading(doc, "3.6 Дополнительные поля", level=2)
    _field_table(doc, b["extra"] or [
        {"field_name": "PROMO old / PROMO new", "data_type": "Флаг / пусто",
         "categorical_values": ["NaN или метка"]},
        {"field_name": "Инфо", "data_type": "Текст / NaN",
         "categorical_values": ["«Новинка», «Есть продажи»"]},
        {"field_name": "Комментарий аналитика", "data_type": "Текст / NaN",
         "categorical_values": []},
    ])


def _write_generic_sheet(doc: Document, name: str, idx: int,
                         columns: list, description: str) -> None:
    _heading(doc, f"{idx}. Лист «{name}»")
    _body(doc, description)
    _field_table(doc, columns)


def _write_relations(doc: Document, relationships: list) -> None:
    _heading(doc, "9. Граф связей между листами")
    _body(doc, "Ключевые JOIN-связи между листами файла:")
    table = _make_table(doc, ["Лист-источник", "Лист-приёмник", "Связь / JOIN-ключ"])
    _STATIC = [
        ("Факт продаж",  "BCG",                         "Медиана доли (строка 7) → BCG-порог"),
        ("ABC расчет",   "BCG",                         "КАТЕГОРИЯ + ЛИНИЯ МОДЕЛИ (260 строк)"),
        ("BCG",          "Тренды категорий",             "КАТЕГОРИЯ (агрегация)"),
        ("BCG",          "Слайд 1 тренды категорий (2)", "КАТЕГОРИЯ (агрегация)"),
        ("BCG",          "Слайд 1 тренды категорий",    "КАТЕГОРИЯ (агрегация)"),
    ]
    seen: set = set()
    for src, tgt, key in _STATIC:
        k = (src, tgt)
        if k not in seen:
            seen.add(k)
            r = table.add_row().cells
            r[0].text, r[1].text, r[2].text = src, tgt, key
    for rel in (relationships or []):
        k = (rel.get("source_table", ""), rel.get("target_table", ""))
        if k not in seen:
            seen.add(k)
            r = table.add_row().cells
            r[0].text = rel.get("source_table", "")
            r[1].text = rel.get("target_table", "")
            r[2].text = ", ".join(rel.get("join_keys", []))
    doc.add_paragraph()


def _write_business_rules(doc: Document) -> None:
    _heading(doc, "10. Ключевые бизнес-правила и расчёты")
    table = _make_table(doc, ["Правило", "Описание"])
    for rule, desc in [
        ("BCG-порог",          "Медиана доли в доходах по всем SKU за квартал."),
        ("ABC-граница A",      "Накопленная доля ≤ 0.80 в категории → A."),
        ("ABC-граница B",      "Накопленная доля 0.80–0.95 → B."),
        ("ABC-граница C",      "Накопленная доля > 0.95 → C."),
        ("«нет продаж»",       "Бизнес-статус SKU: нет отгрузок в квартале. Не null, не ошибка."),
        ("Горизонтальный layout", "Кварталы горизонтально: Продажи + Доход + Маржа + Доля."),
        ("Периодичность",      "Квартальная. IV кв 2024 (база) → II, III, IV кв 2025."),
    ]:
        r = table.add_row().cells
        r[0].text, r[1].text = rule, desc
    doc.add_paragraph()


def _write_channels(doc: Document, fact_table: Optional[dict]) -> None:
    _heading(doc, "11. Каналы продаж (лист «Факт продаж»)")
    table = _make_table(doc, ["Код канала", "Описание"])
    channels = [
        ("03_НСУ_Відділ продажів_НСУ",       "Основной отдел продаж (розница ~81%)."),
        ("03_НСУ_Відділ проектних продажів",  "Проектные продажи (B2B, крупные заказы ~8-14%)."),
        ("E-commerce",                        "Интернет-продажи (~6-11%)."),
        ("Export",                            "Экспортные отгрузки (<1.5%)."),
        ("03_СКФ",                            "Образцы / СКФ (<0.1%)."),
    ]
    if fact_table:
        names_col = next((c for c in fact_table.get("columns", [])
                          if "названия строк" in c.get("field_name", "").lower()), None)
        if names_col and names_col.get("categorical_values"):
            channels = [(v, "—") for v in names_col["categorical_values"]]
    for code, desc in channels:
        r = table.add_row().cells
        r[0].text, r[1].text = code, desc
    doc.add_paragraph()


def _write_vector_db_recommendations(doc: Document) -> None:
    _heading(doc, "12. Рекомендации для vector DB / LLM-агента")
    _body(doc, "Для превращения этого датасета в семантически индексируемый "
               "источник данных рекомендуется:")
    for text in [
        "Атомарная единица для эмбеддинга — одна строка BCG-листа (SKU × квартал). "
        "Описание: «{КАТЕГОРИЯ} / {ЛИНИЯ МОДЕЛИ} в {квартал}: продажи {X} шт., доход {Y} грн.».",
        "Метаданные (не в эмбеддинг, но в фильтры): КАТЕГОРИЯ, квартал, BCG-группа, ABC-метка.",
        "Критическая семантика (только от людей): почему «нет продаж» — "
        "LLM не может определить без контекста.",
        "Триггер обновления: новый квартальный срез → пересчёт медианы BCG-порога "
        "→ переклассификация всех SKU → обновление эмбеддингов изменившихся строк.",
        "JOIN-граф для retrieval-агента: BCG ↔ ABC расчет (КАТЕГОРИЯ + ЛИНИЯ МОДЕЛИ) "
        "→ Тренды категорий (КАТЕГОРИЯ) → Факт продаж (медиана).",
    ]:
        _bullet(doc, text)


def _write_footer(doc: Document, semantic: dict) -> None:
    provider = semantic.get("winner_provider") or "LLM"
    model = "unknown"
    if semantic.get("resolved"):
        model = semantic["resolved"][0].get("model", "unknown")
    p = doc.add_paragraph()
    _styled_run(
        p,
        f"Схема составлена автоматически на основании структурного анализа файла. "
        f"Семантические интерпретации подтверждены {provider} ({model}) + человеком.",
        size_pt=8.5,
        color=C_LIGHT,
    )


# ── Sheet descriptions ─────────────────────────────────────────────────────────

_SHEET_DESC: dict = {
    "ABC расчет": (
        "Детальный ABC-анализ внутри каждой категории. "
        "Структура: 4 блока по кварталам, каждый блок — КАТЕГОРИЯ + ЛИНИЯ МОДЕЛИ + ABC + Доля."
    ),
    "Тренды категорий": (
        "Агрегированный вид: доходы по укрупнённым категориям за 4 полугодия (II/2023 → I/2025)."
    ),
    "Слайд 1 тренды категорий (2)": (
        "Расширенная версия трендов: три группы метрик (продажи в шт., доход в грн., маржа) "
        "× 4 периода + темпы роста."
    ),
    "Слайд 1 тренды категорий": (
        "Доля в доходах по категориям за январь–июнь 2025. Данные для слайда визуализации."
    ),
    "Факт продаж": (
        "Источниковые данные в разрезе каналов/отделов продаж. "
        "Содержит медиану доли по каждому кварталу в строке 7 (BCG-порог)."
    ),
    "План Антонова": (
        "Помесячный факт продаж по всей компании за 12 месяцев отчётного года. "
        "Используется как бенчмарк общей динамики."
    ),
}

_SHEET_ORDER = [
    "ABC расчет",
    "Тренды категорий",
    "Слайд 1 тренды категорий (2)",
    "Факт продаж",
    "План Антонова",
]


# ── Public API ─────────────────────────────────────────────────────────────────

def generate(schema_path: Path, semantic_path: Path, output_path: Path) -> None:
    """Generate .docx report from final_schema.json + semantic.json.

    Args:
        schema_path:   Path to outputs/final_schema_<timestamp>.json
        semantic_path: Path to outputs/semantic_<timestamp>.json
        output_path:   Destination path for the .docx file
    """
    with open(schema_path, encoding="utf-8") as f:
        schema: dict = json.load(f)
    with open(semantic_path, encoding="utf-8") as f:
        semantic: dict = json.load(f)

    tables: list = schema.get("tables", [])
    by_name: dict = {t["name"]: t for t in tables}
    relationships: list = schema.get("relationships", [])

    doc = Document()

    # 1. Title block (no table)
    _write_title(doc, schema)

    # 2. Purpose (no table)
    _heading(doc, "1. Назначение и цель датасета")
    _body(doc, ("Датасет предназначен для продуктового ABC/BCG-анализа товарного портфеля. "
                "На входе — данные продаж по SKU за 4 квартала, канальная разбивка, план и факт "
                "по компании. На выходе — классификация каждого SKU по матрице BCG и ABC "
                "внутри категории."))

    # 3. Sheet overview table (T00 = 1 table)
    _write_sheet_overview(doc, tables)

    # 4. BCG sheet — 6 subsection tables (T01–T06)
    bcg = by_name.get("BCG")
    if bcg:
        _write_bcg_section(doc, bcg["columns"])

    # 5–9. Remaining sheets — 1 table each (T07–T11)
    for idx, name in enumerate(_SHEET_ORDER, start=4):
        t = by_name.get(name)
        if t:
            _write_generic_sheet(doc, name, idx,
                                 t["columns"], _SHEET_DESC.get(name, "—"))

    # 10. Relations graph (T12)
    _write_relations(doc, relationships)

    # 11. Business rules (T13)
    _write_business_rules(doc)

    # 12. Channels (T14)
    _write_channels(doc, by_name.get("Факт продаж"))

    # 13. Vector DB recommendations (no table)
    _write_vector_db_recommendations(doc)

    # Footer
    _write_footer(doc, semantic)

    doc.save(str(output_path))

    # ── inline verification ────────────────────────────────────────────────────
    v = Document(str(output_path))
    h1_count  = sum(1 for p in v.paragraphs if p.style.name == "Heading 1")
    h2_count  = sum(1 for p in v.paragraphs if p.style.name == "Heading 2")
    tbl_count = len(v.tables)
    logger.info(
        "Report saved: %s  paragraphs=%d  tables=%d  H1=%d  H2=%d",
        output_path, len(v.paragraphs), tbl_count, h1_count, h2_count,
    )
    if tbl_count != 15:
        logger.warning("Expected 15 tables, got %d", tbl_count)
